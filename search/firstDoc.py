from docx import Document
from datetime import datetime
from template.currentTemplate import *
import re
import zipfile
import os
import glob


def replaceArgs(input, **args):
	for key, value in args.iteritems():
		input = input.replace('['+str(key).strip()+']', str(value))
	return input

def getResume(resumePath):
	if resumePath is None:
		return "Resume not Found"
	else:
		resume=''
		docR = Document(resumePath)
		tables = docR.tables
		for table in tables:
			for row in table.rows:
				resume += "<ul>"
				for cell in row.cells:
					for para in cell.paragraphs:
						resume += "<li>" + para.text + "</li>"
		resume += "</ul>"

		for para in docR.paragraphs:
			resume += "<br/>"
			resume +=para.text
		return resume	
	
def getCoverLetter(coverLetterPath):
	if coverLetterPath is None:
		return "Cover Letter not Found"
	else:
		coverLetter=''
		docC = Document(coverLetterPath)
		for para in docC.paragraphs:
			coverLetter +=para.text
			coverLetter += "<br/>"
		return coverLetter
							
def submit(request, **args):
	files = glob.glob('./empty/*')
	for f in files:
		os.remove(f)

	for a in args.keys():
		print a
	docR = Document(getCurrentResume(request))
	tables = docR.tables
	for table in tables:
	    for row in table.rows:
	        for cell in row.cells:
	            for paragraph in cell.paragraphs:
					print paragraph.text
					paragraph.text = replaceArgs(paragraph.text, **args);
					print paragraph.text
	for para in docR.paragraphs:
		for key in args.keys():
			if key in para.text:
				para.text = replaceArgs(para.text, **args);
	resumeSavePath = re.sub('documents','empty',str(getCurrentResume(request)))
	docR.save(resumeSavePath)
	
	docC = Document(getCurrentCoverLetter(request))
	for para in docC.paragraphs:	
		for key in args.keys():
			if key in para.text:
				para.text = replaceArgs(para.text, **args);
	coverLetterSavePath = re.sub('documents','empty',str(getCurrentCoverLetter(request)))
	docC.save(coverLetterSavePath)
	
	emailBody = replaceArgs(getCurrentEmailBody(request), **args)
	emailSubject = replaceArgs(getCurrentEmailSubject(request), **args)
	emailDocument = Document()
	emailDocument.add_heading('Contents of Application Email')
	emailDocument.add_paragraph('Email Subject:')
	emailDocument.add_paragraph(emailSubject)
	emailDocument.add_paragraph('Email Body:')
	emailDocument.add_paragraph(emailBody)
	emailDocument.save("./empty/emailDocument.docx")


	zipfileName = datetime.now().strftime("%Y%m%d") + getCurrentTemplateName(request) + ".zip"
	zf = zipfile.ZipFile(zipfileName, "w")
	zf.write(resumeSavePath)
	zf.write(coverLetterSavePath)
	zf.write("./empty/emailDocument.docx")
	zf.close()
	os.rename(("./"+zipfileName), ("./empty/"+zipfileName))

	return "./empty/" + zipfileName
	
