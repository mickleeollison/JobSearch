from template.models import *
from login.models import User
from docx import Document
import re


def getCurrentResume(request):
	try:
		t = Template.objects.get(id = request.session['templateId'])
		return t.resume
	except:
		return None

def getCurrentTemplateName(request):
	try:
		t = Template.objects.get(id = request.session['templateId'])
		return t.name
	except:
		return None

def getCurrentCoverLetter(request):
	try:
		t = Template.objects.get(id = request.session['templateId'])
		return t.coverLetter
	except:
		return None

def getCurrentEmailBody(request):
	try:
		t = Template.objects.get(id = request.session['templateId'])
		return t.emailBody
	except:
		return "No Current Template Selected"

def getCurrentEmailSubject(request):
	try:
		t = Template.objects.get(id = request.session['templateId'])
		return t.emailSubject
	except:
		return "No Current Template Selected"

def getCurrentEmailName(request):
	try:
		t = Template.objects.get(id = request.session['templateId'])
		return t.name
	except:
		return "No Current Template Selected"

def getTemplates(request):
	ret = []
	user = User.objects.filter(email = request.session['email'])[0]
	templates = Template.objects.filter(user= user)
	print templates
	for t in templates:
		temp = {}
		temp['id'] = int(t.id)
		temp['name'] = t.name
		temp['resume'] = t.resume
		temp['emailSubject'] = t.emailSubject
		temp['emailBody'] = t.emailBody
		temp['coverLetter'] = t.coverLetter
		ret.append(temp)
	return ret

def getTemplate(request):
	try:
		plugins = []
		t = Template.objects.get(id = request.session['templateId'])
		matchObj = re.search(r'\[.*\]',t.emailSubject)
		if matchObj:
			words = matchObj.group(0).split(" ")
			for word in words:
				if "[" in word:
					plugin = re.sub(r'\[|\]|[^\w\s]', '', word)
					if plugin not in plugins:
						plugins.append(plugin)
		matchObj = re.search(r'\[.*\]',re.sub(r"\n+", " ", t.emailBody))
		if matchObj:
			words = matchObj.group(0).split(" ")
			for word in words:
				if "[" in word:
					plugin = re.sub(r'\[|\]|[^\w\s]', '', word)
					if plugin not in plugins:
						plugins.append(plugin)
		getPlugins(getCurrentResume(request),getCurrentCoverLetter(request),plugins)
		return plugins
	except BaseException, e:
		print e

def getPlugins(resumePath,coverLetterPath, plugins):
	docR = Document(resumePath)
	tables = docR.tables
	for table in tables:
	    for row in table.rows:
	        for cell in row.cells:
	            for paragraph in cell.paragraphs:
					matchObj = re.search(r'\[.*\]',paragraph.text)
					if matchObj:
						words = matchObj.group(0).split(" ")
						for word in words:
							if "[" in word:
								plugin = re.sub(r'\[|\]|[^\w\s]', '', word)
								if plugin not in plugins:
									plugins.append(plugin)
						
	for paragraph in docR.paragraphs:
		matchObj = re.search(r'\[.*\]',paragraph.text)
		if matchObj:
			words = matchObj.group(0).split(" ")
			for word in words:
				if "[" in word:
					plugin = re.sub(r'\[|\]|[^\w\s]', '', word)
					if plugin not in plugins:
						plugins.append(plugin)
	docC = Document(coverLetterPath)
	for paragraph in docC.paragraphs:
		matchObj = re.search(r'\[.*\]',paragraph.text)
		if matchObj:
			words = matchObj.group(0).split(" ")
			for word in words:
				if "[" in word:
					plugin = re.sub(r'\[|\]|[^\w\s]', '', word)
					if plugin not in plugins:
						plugins.append(plugin)
	return plugins
